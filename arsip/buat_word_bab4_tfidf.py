"""
Script untuk membuat dokumen Word BAB 4 sub-bab 2 (atau 4.x): TF-IDF
Sesuai referensi Subbab_TF_IDF_Bab4.docx dan Colab output.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "BAB4_4_TF_IDF.docx") # Bisa disesuaikan jadi 4.2 atau 4.4

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)
LINE_SPACING = 1.5

def set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, 
                         space_after=Pt(0), space_before=Pt(0), 
                         first_line_indent=Cm(1.27), line_spacing=LINE_SPACING):
    pf = paragraph.paragraph_format
    pf.alignment = alignment
    pf.space_after = space_after
    pf.space_before = space_before
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent

def add_run(paragraph, text, bold=False, italic=False, color=None, font_name=FONT_NAME, font_size=FONT_SIZE):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    run.font.name = font_name
    run.font.size = font_size
    # Set East Asian font
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rPr.insert(0, rFonts)
    return run

def add_normal_paragraph(doc, text_parts, first_line_indent=Cm(1.27)):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_indent=first_line_indent)
    for part in text_parts:
        if len(part) == 3:
            text, bold, italic = part
            color = None
        elif len(part) == 4:
            text, bold, italic, color = part
        add_run(p, text, bold=bold, italic=italic, color=color)
    return p

def add_formula_line(doc, formula_text, italic=True):
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, 
                         first_line_indent=None, space_before=Pt(6), space_after=Pt(6))
    add_run(p, formula_text, italic=italic)
    return p

def add_screenshot_placeholder(doc, text):
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, 
                         first_line_indent=None, space_before=Pt(12), space_after=Pt(12))
    add_run(p, text, bold=True, color=RGBColor(255, 0, 0))
    
    # Keterangan gambar
    p_ket = doc.add_paragraph()
    set_paragraph_format(p_ket, alignment=WD_ALIGN_PARAGRAPH.CENTER, 
                         first_line_indent=None, space_before=Pt(0), space_after=Pt(12))
    add_run(p_ket, "Gambar 4.x [Beri Keterangan Gambar Di Sini]", bold=False, italic=False)

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = FONT_NAME
font.size = FONT_SIZE

# Set margin
for section in doc.sections:
    section.top_margin = Cm(4)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

# Heading
p_heading = doc.add_paragraph()
set_paragraph_format(p_heading, alignment=WD_ALIGN_PARAGRAPH.LEFT, 
                     first_line_indent=None, space_before=Pt(12), space_after=Pt(6))
add_run(p_heading, "4. Pembobotan ", bold=True)
add_run(p_heading, "Term Frequency - Inverse Document Frequency", bold=True, italic=True)
add_run(p_heading, " (TF-IDF)", bold=True)

# Paragraf 1
add_normal_paragraph(doc, [
    ("Setelah data ulasan melewati tahap ", False, False),
    ("preprocessing", False, True),
    (", data yang semula berbentuk teks perlu diubah menjadi bentuk numerik agar dapat diproses oleh algoritma klasifikasi. Pada penelitian ini, proses perubahan teks menjadi nilai numerik dilakukan menggunakan metode ", False, False),
    ("Term Frequency-Inverse Document Frequency", False, True),
    (" (TF-IDF). Metode ini digunakan untuk memberikan bobot pada setiap kata berdasarkan tingkat kemunculan kata dalam dokumen dan tingkat kepentingan kata tersebut pada keseluruhan dataset ulasan mobil listrik Wuling Air EV.", False, False)
])

# Paragraf 2
add_normal_paragraph(doc, [
    ("Data yang digunakan pada tahap TF-IDF merupakan hasil dari proses ", False, False),
    ("case folding, cleaning, tokenizing", False, True),
    (", normalisasi kata, ", False, False),
    ("stopword removal", False, True),
    (" dan ", False, False),
    ("stemming", False, True),
    (". Setelah proses tersebut selesai, setiap ulasan menjadi lebih bersih dan memiliki susunan kata yang lebih siap untuk dihitung bobotnya. Dengan demikian, kata-kata yang tidak relevan telah dikurangi, sedangkan kata yang memiliki informasi penting terhadap sentimen tetap dipertahankan untuk digunakan pada tahap pembobotan.", False, False)
])

# Paragraf 3
add_normal_paragraph(doc, [
    ("Pada tahap ini, setiap ulasan diperlakukan sebagai satu dokumen. Sistem menghitung bobot kata pada masing-masing dokumen agar kata yang memiliki pengaruh lebih besar terhadap isi ulasan dapat diberi nilai yang lebih tinggi. Kata yang sering muncul dalam sebuah ulasan akan memiliki nilai ", False, False),
    ("Term Frequency", False, True),
    (" (TF) yang tinggi. Namun, kata yang terlalu sering muncul pada banyak dokumen tidak selalu dianggap penting, sehingga perlu dihitung kembali menggunakan nilai ", False, False),
    ("Inverse Document Frequency", False, True),
    (" (IDF). Kombinasi kedua perhitungan tersebut menghasilkan bobot TF-IDF yang digunakan sebagai fitur masukan bagi model klasifikasi.", False, False)
])

# Paragraf 4
add_normal_paragraph(doc, [
    ("Rumus perhitungan bobot TF-IDF yang digunakan dapat dituliskan sebagai berikut:", False, False)
])

add_formula_line(doc, "TF-IDF(t, d) = TF(t, d) × IDF(t)")
add_formula_line(doc, "TF(t, d) = 1 + log(freq(t, d))")
add_formula_line(doc, "IDF(t) = log(N / df(t)) + 1")

# Paragraf 5
add_normal_paragraph(doc, [
    ("Keterangan dari rumus tersebut adalah ", False, False),
    ("TF-IDF(t,d)", False, True),
    (" sebagai bobot akhir suatu term ", False, False),
    ("t", False, True),
    (" pada dokumen ", False, False),
    ("d", False, True),
    (", ", False, False),
    ("TF(t,d)", False, True),
    (" sebagai frekuensi kemunculan term ", False, False),
    ("t", False, True),
    (" pada dokumen ", False, False),
    ("d", False, True),
    (", ", False, False),
    ("IDF(t)", False, True),
    (" sebagai nilai yang menunjukkan tingkat keunikan term terhadap seluruh dokumen, ", False, False),
    ("N", False, True),
    (" sebagai jumlah seluruh dokumen, dan ", False, False),
    ("df(t)", False, True),
    (" sebagai jumlah dokumen yang memuat term tersebut.", False, False)
])

# Paragraf 6
add_normal_paragraph(doc, [
    ("Dalam implementasi program, proses pembobotan TF-IDF dikonfigurasi dengan parameter ", False, False),
    ("ngram_range", False, True),
    ("=(1,2) untuk menangkap fitur unigram dan bigram, ", False, False),
    ("min_df", False, True),
    ("=2 untuk mengabaikan kata yang terlalu jarang muncul, dan ", False, False),
    ("max_df", False, True),
    ("=0,95 untuk mengabaikan kata yang terlalu umum. Proses ini dilakukan setelah data ulasan selesai dibersihkan. Setiap kata hasil ", False, False),
    ("preprocessing", False, True),
    (" diubah menjadi fitur numerik yang merepresentasikan bobot kata dalam dataset. Berikut ini adalah hasil perhitungan bobot TF-IDF pada beberapa sampel ulasan:", False, False)
])

# Placeholder Screenshot 1
add_screenshot_placeholder(doc, "[INSERT SCREENSHOT CELL 8 DI SINI: TABEL PEMBOBOTAN TF-IDF 15 SAMPLE DATA]")

# Paragraf 7
add_normal_paragraph(doc, [
    ("Berdasarkan contoh pada tabel di atas, kata-kata yang lebih sering muncul atau memiliki bobot makna khusus akan mendapatkan skor TF-IDF yang lebih tinggi dalam masing-masing dokumen. Selanjutnya, dari keseluruhan dokumen juga dapat dihitung rata-rata bobot TF-IDF tertinggi untuk mengetahui kata atau fitur mana yang memiliki pengaruh paling signifikan secara global. Berikut adalah gabungan pembobotan untuk fitur-fitur tertinggi pada sampel dokumen:", False, False)
])

# Placeholder Screenshot 2
add_screenshot_placeholder(doc, "[INSERT SCREENSHOT CELL 9 DI SINI: TABEL GABUNGAN PEMBOBOTAN TF-IDF TOP 15 FITUR]")

# Placeholder Screenshot 3
add_screenshot_placeholder(doc, "[INSERT SCREENSHOT CELL 10 DI SINI: STATISTIK KESELURUHAN DAN TOP 20 KATA]")

# Paragraf 8
add_normal_paragraph(doc, [
    ("Sebagai contoh, kata yang sering muncul pada satu ulasan tetapi jarang muncul pada keseluruhan dataset akan memperoleh bobot yang lebih tinggi karena dianggap lebih mewakili isi dokumen tersebut. Sebaliknya, kata yang muncul hampir di seluruh dokumen akan memiliki bobot yang lebih rendah karena dianggap kurang mampu membedakan antar ulasan. Dengan cara ini, proses TF-IDF membantu sistem dalam memilih fitur kata yang lebih relevan sebelum data diproses oleh model SVM.", False, False)
])

# Paragraf 9
add_normal_paragraph(doc, [
    ("Hasil akhir dari tahap TF-IDF adalah data numerik yang siap digunakan pada subbab klasifikasi ", False, False),
    ("Support Vector Machine", False, True),
    (". Dengan adanya pembobotan ini, model SVM dapat mempelajari pola perbedaan antar kelas sentimen berdasarkan bobot kata yang terbentuk. Oleh karena itu, tahap TF-IDF menjadi penghubung penting antara proses ", False, False),
    ("preprocessing", False, True),
    (" dan proses klasifikasi, karena tanpa tahap ini data teks tidak dapat langsung diproses oleh algoritma ", False, False),
    ("machine learning", False, True),
    (".", False, False)
])

doc.save(OUTPUT_FILE)
print(f"[OK] Dokumen berhasil dibuat: {OUTPUT_FILE}")
